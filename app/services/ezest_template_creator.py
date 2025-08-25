import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

class EZestTemplateCreator:
    """Creates e-Zest formatted resume templates"""

    def __init__(self, templates_dir: Path):
        self.templates_dir = templates_dir
        self.doc = Document()
        self._set_document_margins()
        self._create_styles()

    def _set_document_margins(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def _create_styles(self):
        styles = self.doc.styles
        
        # Normal style
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading styles
        for i in range(1, 3):
            heading_style = styles[f'Heading {i}']
            font = heading_style.font
            font.name = 'Calibri'
            font.bold = True
            if i == 1:
                font.size = Pt(16)
                font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
            elif i == 2:
                font.size = Pt(14)
                font.color.rgb = RGBColor(0x17, 0x36, 0x5D)

    def _add_section_heading(self, text):
        heading = self.doc.add_paragraph(text, style='Heading 2')
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

    def create_ezest_template(self) -> Path:
        """Create a professional e-Zest style template"""
        
        # Header
        name_paragraph = self.doc.add_paragraph()
        name_paragraph.add_run('{{contact_info.name}}').bold = True
        name_paragraph.style.font.size = Pt(16)
        name_paragraph.style.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        
        # Image - This will be tricky with docxtpl, let's add a placeholder
        # self.doc.add_picture('path_to_image.png', width=Inches(6.5))
        self.doc.add_paragraph("{{title}}")


        # Profile Summary
        self._add_section_heading('Profile Summary')
        summary_table = self.doc.add_table(rows=1, cols=1)
        summary_table.style = 'Table Grid'
        summary_cell = summary_table.cell(0, 0)
        summary_cell.text = '{{summary}}'

        # Professional Skills
        self._add_section_heading('Professional Skills')
        skills_table = self.doc.add_table(rows=1, cols=2)
        skills_table.style = 'Table Grid'
        skills_table.cell(0,0).text = "Category"
        skills_table.cell(0,1).text = "Skills"
        skills_table.add_row().cells[0].text = "{% for skill in skills %}{{ skill.category }}{% endfor %}"
        skills_table.add_row().cells[1].text = "{% for skill in skills %}{{ skill.skills|join(', ') }}{% endfor %}"


        # Relevant Work Experience
        self._add_section_heading('Relevant Work Experience')
        exp_table = self.doc.add_table(rows=1, cols=1)
        exp_table.style = 'Table Grid'
        exp_table.cell(0,0).text = "{% for exp in experience %}"
        exp_table.add_row().cells[0].text = "Project #{{ loop.index }}"
        exp_table.add_row().cells[0].text = "Duration: {{ exp.start_date }} - {{ exp.end_date }}"
        exp_table.add_row().cells[0].text = "Project Description: {{ exp.description }}"
        exp_table.add_row().cells[0].text = "Technology: {{ exp.technologies|join(', ') }}"
        exp_table.add_row().cells[0].text = "Role & Responsibilities: {{ exp.responsibilities|join(', ') }}"
        exp_table.add_row().cells[0].text = "{% endfor %}"

        # Other Notable Projects
        self._add_section_heading('Other Notable Projects')
        other_projects_table = self.doc.add_table(rows=1, cols=3)
        other_projects_table.style = 'Table Grid'
        other_projects_table.cell(0,0).text = "Project"
        other_projects_table.cell(0,1).text = "Duration"
        other_projects_table.cell(0,2).text = "Technology"
        other_projects_table.add_row().cells[0].text = "{% for project in other_projects %}{{ project.name }}{% endfor %}"
        other_projects_table.add_row().cells[1].text = "{% for project in other_projects %}{{ project.duration }}{% endfor %}"
        other_projects_table.add_row().cells[2].text = "{% for project in other_projects %}{{ project.technologies|join(', ') }}{% endfor %}"

        # Education Details
        self._add_section_heading('Education Details')
        edu_table = self.doc.add_table(rows=1, cols=3)
        edu_table.style = 'Table Grid'
        edu_table.cell(0,0).text = "Course"
        edu_table.cell(0,1).text = "University"
        edu_table.cell(0,2).text = "Year of Passing"
        edu_table.add_row().cells[0].text = "{% for edu in education %}{{ edu.degree }}{% endfor %}"
        edu_table.add_row().cells[1].text = "{% for edu in education %}{{ edu.institution }}{% endfor %}"
        edu_table.add_row().cells[2].text = "{% for edu in education %}{{ edu.graduation_year }}{% endfor %}"


        template_path = self.templates_dir / "ezest.docx"
        self.doc.save(template_path)

        return template_path