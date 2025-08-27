import os
import uuid
import re
from pathlib import Path
from typing import Optional, Dict, Any
import asyncio
from docx import Document
try:
    import docx2txt
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False
import aiofiles
import logging

from app.core.config import settings
from app.models.schemas import ExtractedData, ContactInfo, Experience, Education
from app.services.nlp_extractor import NLPExtractor
from app.services.ocr_processor import OCRProcessor
from app.services.gemini_processor import GeminiProcessor
import fitz  # PyMuPDF

class DocumentProcessor:
    """Core document processing service"""
    
    def __init__(self, gemini_api_key: Optional[str] = None):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.output_dir = Path(settings.OUTPUT_DIR)
        self.nlp_extractor = NLPExtractor()
        self.ocr_processor = OCRProcessor()
        self.gemini_processor = None
        if gemini_api_key:
            try:
                self.gemini_processor = GeminiProcessor(gemini_api_key)
            except Exception as e:
                logging.warning(f"Failed to initialize Gemini processor: {e}")
        
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return unique filename"""
        file_extension = Path(filename).suffix.lower()
        if file_extension not in settings.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
        # Generate unique filename
        unique_id = str(uuid.uuid4())
        unique_filename = f"{unique_id}_{filename}"
        file_path = self.upload_dir / unique_filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
            
        return unique_filename
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract raw text from DOCX file using docx2txt for better extraction"""
        try:
            # Use docx2txt for better text extraction including tables if available
            if HAS_DOCX2TXT:
                text = docx2txt.process(str(file_path))
                if text and text.strip():
                    return text.strip()
            
            # Fallback to python-docx
            doc = Document(file_path)
            text_content = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
                            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF with block-ordering similar to reference implementation"""
        try:
            text = ""
            with fitz.open(str(file_path)) as doc:
                for page in doc:
                    blocks = page.get_text("blocks")
                    # Sort by y (top) then x (left)
                    blocks.sort(key=lambda b: (b[1], b[0]))
                    text += "\n".join(b[4] for b in blocks)
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def advanced_data_extraction(self, text: str) -> ExtractedData:
        """Advanced data extraction using NLP and pattern matching"""
        try:
            # Use NLP extractor for comprehensive data extraction
            contact_info = self.nlp_extractor.extract_contact_info(text)
            experience = self.nlp_extractor.extract_experience(text)
            education = self.nlp_extractor.extract_education(text)
            skills = self.nlp_extractor.extract_skills(text)
            
            # Extract summary/objective section
            summary = self._extract_summary(text)
            
            # Calculate confidence score based on extracted data completeness
            confidence_score = self._calculate_confidence_score(
                contact_info, experience, education, skills
            )
            
            return ExtractedData(
                contact_info=contact_info,
                summary=summary,
                experience=experience,
                education=education,
                skills=skills,
                raw_text=text,
                confidence_score=confidence_score
            )
            
        except Exception as e:
            logging.error(f"Advanced extraction failed: {str(e)}")
            # Fallback to basic extraction
            return self.basic_data_extraction(text)
    
    def basic_data_extraction(self, text: str) -> ExtractedData:
        """Basic data extraction using simple heuristics (fallback method)"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Initialize extracted data
        contact_info = ContactInfo()
        experience = []
        education = []
        skills = []
        
        # Simple email detection
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info.email = emails[0]
        
        # Simple phone detection
        phone_pattern = r'(\+?1?[-.\.\s]?)?\(?([0-9]{3})\)?[-.\.\s]?([0-9]{3})[-.\.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info.phone = ''.join(phones[0])
        
        # Simple name extraction (assume first non-empty line is name)
        if lines:
            potential_name = lines[0]
            # Basic validation - name shouldn't contain @ or numbers
            if '@' not in potential_name and not any(char.isdigit() for char in potential_name):
                contact_info.name = potential_name
        
        # Heuristic professional title extraction (best-effort)
        # Look at the next few non-empty lines for a short role/title phrase
        common_titles = [
            'developer','engineer','administrator','consultant','analyst','architect',
            'manager','lead','specialist','scientist','designer','tester','qa','devops',
            'sdet','full stack','frontend','backend','cloud','data','ml','ai','security'
        ]
        for i in range(1, min(6, len(lines))):
            line = lines[i].strip()
            if not line or len(line) > 70:
                continue
            low = line.lower()
            if any(t in low for t in common_titles):
                # Remove stray parentheses and bullets
                cleaned = line.strip(' ()•-\t')
                # Avoid mistaking email/phone or headings
                if '@' in cleaned or any(ch.isdigit() for ch in cleaned):
                    continue
                contact_info.title = cleaned
                break
        
        # Extract skills section
        skills_keywords = ['skills', 'technical skills', 'core competencies', 'technologies']
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in skills_keywords):
                # Look for skills in next few lines
                for j in range(i+1, min(i+10, len(lines))):
                    if lines[j] and not any(stop_word in lines[j].lower() 
                                          for stop_word in ['experience', 'education', 'work']):
                        # Split by common delimiters
                        line_skills = re.split(r'[,;|•·]', lines[j])
                        skills.extend([skill.strip() for skill in line_skills if skill.strip()])
                break
        
        return ExtractedData(
            contact_info=contact_info,
            experience=experience,
            education=education,
            skills=skills[:10],  # Limit to first 10 skills
            raw_text=text,
            confidence_score=0.5  # Lower confidence for basic extraction
        )
    
    async def process_document(self, filename: str, use_gemini: bool = False, gemini_api_key: Optional[str] = None) -> ExtractedData:
        """Process uploaded document and extract data"""
        file_path = self.upload_dir / filename
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {filename}")
        
        # Extract text based on file type
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
            # If text is too sparse, consider OCR fallback later
            if not text or len(text.strip()) < 20:
                # TODO: Add optional OCR fallback for scanned PDFs
                pass
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Check if Gemini should be used
        if use_gemini and (gemini_api_key or self.gemini_processor):
            try:
                # Update API key if provided
                if gemini_api_key and not self.gemini_processor:
                    self.gemini_processor = GeminiProcessor(gemini_api_key)
                elif gemini_api_key and self.gemini_processor:
                    self.gemini_processor.update_api_key(gemini_api_key)
                
                # Use Gemini for extraction
                logging.info("Using Gemini for resume extraction")
                extracted_data = await self.gemini_processor.process_resume(text)
            except Exception as e:
                logging.error(f"Gemini extraction failed: {e}")
                # Fallback to traditional extraction
                extracted_data = self.advanced_data_extraction(text)
        else:
            # Extract structured data using advanced NLP methods
            extracted_data = self.advanced_data_extraction(text)
        
        return extracted_data
    
    def cleanup_file(self, filename: str) -> bool:
        """Clean up uploaded file"""
        try:
            file_path = self.upload_dir / filename
            if file_path.exists():
                file_path.unlink()
                return True
        except Exception:
            pass
        return False
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract summary/objective section from resume"""
        summary_patterns = [
            r'(?i)(?:professional\s+)?summary',
            r'(?i)profile\s+summary',
            r'(?i)objective',
            r'(?i)profile',
            r'(?i)about\s+me',
            r'(?i)career\s+summary'
        ]
        
        lines = text.split('\n')
        summary_start = -1
        
        # Find summary section start
        for i, line in enumerate(lines):
            for pattern in summary_patterns:
                if re.search(pattern, line) and len(line.strip()) < 80:
                    summary_start = i
                    break
            if summary_start != -1:
                break
        
        if summary_start == -1:
            return None
        
        # Extract summary content (next few lines until next section)
        summary_lines = []
        section_end_keywords = [
            'experience', 'education', 'skills', 'technical skills', 
            'work history', 'employment', 'certifications', 'projects'
        ]
        
        for i in range(summary_start + 1, min(summary_start + 15, len(lines))):
            if i >= len(lines):
                break
                
            line = lines[i].strip()
            if not line:
                continue
            
            # Stop if we hit another section header
            if any(keyword in line.lower() for keyword in section_end_keywords):
                # Check if this looks like a section header (short line, all caps or title case)
                if len(line) < 50 and (line.isupper() or line.istitle()):
                    break
                    
            summary_lines.append(line)
        
        if summary_lines:
            return ' '.join(summary_lines)
        
        return None
    
    def _calculate_confidence_score(self, contact_info: ContactInfo, 
                                   experience: list, education: list, 
                                   skills: list) -> float:
        """Calculate confidence score based on extracted data completeness"""
        score = 0.0
        
        # Contact info scoring (40% of total)
        if contact_info.name:
            score += 0.15
        if contact_info.email:
            score += 0.15
        if contact_info.phone:
            score += 0.10
        
        # Experience scoring (30% of total)
        if experience:
            score += 0.20
            if len(experience) > 1:
                score += 0.10
        
        # Education scoring (15% of total)
        if education:
            score += 0.15
        
        # Skills scoring (15% of total)
        if skills:
            score += 0.10
            if len(skills) > 3:
                score += 0.05
        
        return min(score, 1.0)  # Cap at 1.0
