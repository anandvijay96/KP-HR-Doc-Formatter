from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
import logging

class WorkingEZestTemplateCreator:
    """Creates a working e-Zest template with proper docxtpl syntax"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def create_template(self, template_path: Path):
        """Create working e-Zest template with proper docxtpl/Jinja2 syntax"""
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # Header - Name and Title
            self._add_header(doc)
            
            # Profile Summary Section
            self._add_profile_summary(doc)
            
            # Professional Skills Section
            self._add_professional_skills(doc)
            
            # Relevant Work Experience Section
            self._add_work_experience(doc)
            
            # Other Notable Projects Section
            self._add_notable_projects(doc)
            
            # Education Details Section
            self._add_education_details(doc)
            
            # Save template
            doc.save(template_path)
            self.logger.info(f"Working e-Zest template created at {template_path}")
            
        except Exception as e:
            self.logger.error(f"Error creating working e-Zest template: {str(e)}")
            raise
    
    def _add_header(self, doc):
        """Add header with name and title"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('{{ contact_info.name }}')
        name_run.font.size = Pt(16)
        name_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)  # #17365D
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('({{ contact_info.title if contact_info.title else "Professional Title" }})')
        title_run.font.size = Pt(10)
        title_run.font.name = 'Segoe UI'
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add some spacing
        doc.add_paragraph()
    
    def _add_profile_summary(self, doc):
        """Add Profile Summary section with bulleted table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Profile Summary')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create table for summary
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table borders
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        cell = table.cell(0, 0)
        
        # Add summary text
        summary_para = cell.paragraphs[0]
        summary_run = summary_para.add_run('{{ summary if summary else "Professional summary with key achievements and expertise." }}')
        summary_run.font.size = Pt(10)
        summary_run.font.name = 'Segoe UI'
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
    
    def _add_professional_skills(self, doc):
        """Add Professional Skills section with 2-column table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Professional Skills')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create 2-column table for skills (32% / 67% split)
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(2.0)  # ~32%
        table.columns[1].width = Inches(4.2)  # ~67%
        
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        # Skill categories and values
        skill_categories = [
            ('Technical Skills', '{{ skills[:5] | join(", ") if skills else "Technical skills" }}'),
            ('Programming Languages', '{{ skills[5:10] | join(", ") if skills and skills|length > 5 else "Programming languages" }}'),
            ('Frameworks & Tools', '{{ skills[10:15] | join(", ") if skills and skills|length > 10 else "Frameworks and tools" }}'),
            ('Databases', '{{ skills[15:20] | join(", ") if skills and skills|length > 15 else "Database technologies" }}')
        ]
        
        for i, (category, skills_text) in enumerate(skill_categories):
            # Category cell
            cat_cell = table.cell(i, 0)
            cat_para = cat_cell.paragraphs[0]
            cat_run = cat_para.add_run(category)
            cat_run.font.size = Pt(10)
            cat_run.font.name = 'Segoe UI'
            cat_run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
            cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Skills cell
            skills_cell = table.cell(i, 1)
            skills_para = skills_cell.paragraphs[0]
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.size = Pt(10)
            skills_run.font.name = 'Segoe UI'
            skills_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            skills_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def _add_work_experience(self, doc):
        """Add Relevant Work Experience section"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Relevant Work Experience')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Experience loop template
        exp_para = doc.add_paragraph()
        exp_run = exp_para.add_run('''{% for exp in experience %}
Project: {{ exp.company }} - {{ exp.position }}                                                                                             Duration: {{ exp.start_date }} - {{ exp.end_date }}

Project Description: {{ exp.description if exp.description else "Project description" }}

Technology: {{ exp.technologies | join(", ") if exp.technologies else "Technologies used" }}

Role & Responsibilities:
• {{ exp.responsibilities[0] if exp.responsibilities else "Key responsibility 1" }}
• {{ exp.responsibilities[1] if exp.responsibilities and exp.responsibilities|length > 1 else "Key responsibility 2" }}

{% endfor %}''')
        exp_run.font.size = Pt(10)
        exp_run.font.name = 'Segoe UI'
        exp_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        doc.add_paragraph()
    
    def _add_notable_projects(self, doc):
        """Add Other Notable Projects section with 3-column table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Other Notable Projects')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create 3-column table
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        # Header row with blue background
        header_cells = ['Project', 'Duration', 'Technology']
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            self._set_cell_background(cell, RGBColor(0x17, 0x36, 0x5D))
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row in range(1, 3):
            for col in range(3):
                cell = table.cell(row, col)
                para = cell.paragraphs[0]
                field_map = ["name", "duration", "technology"]
                run = para.add_run(f'{{{{ notable_projects[{row-1}].{field_map[col]} if notable_projects and notable_projects|length > {row-1} else "Data" }}}}')
                run.font.size = Pt(10)
                run.font.name = 'Segoe UI'
                run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def _add_education_details(self, doc):
        """Add Education Details section with 3-column table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Education Details')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create 3-column table
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        # Header row with blue background
        header_cells = ['Course', 'University', 'Year of Passing']
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            self._set_cell_background(cell, RGBColor(0x17, 0x36, 0x5D))
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row in range(1, 3):
            for col in range(3):
                cell = table.cell(row, col)
                para = cell.paragraphs[0]
                field_map = ["degree", "institution", "graduation_date"]
                run = para.add_run(f'{{{{ education[{row-1}].{field_map[col]} if education and education|length > {row-1} else "Data" }}}}')
                run.font.size = Pt(10)
                run.font.name = 'Segoe UI'
                run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def _set_table_borders(self, table, color):
        """Set table borders with specified color"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Use lxml find with qualified names instead of non-existent helpers
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = tcBorders.find(qn(f"w:{border_name}"))
                    if border is None:
                        border = OxmlElement(f'w:{border_name}')
                        tcBorders.append(border)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:color'), f'{color.rgb:06x}')
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:fill'), f'{color.rgb:06x}')
