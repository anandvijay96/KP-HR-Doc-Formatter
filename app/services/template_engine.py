import os
import uuid
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docxtpl import DocxTemplate
import json
from datetime import datetime
import logging
import re
import shutil
import importlib.util
from importlib.machinery import SourceFileLoader

from app.core.config import settings
from app.models.schemas import ExtractedData, TemplateInfo
from app.services.ezest_template_creator import EZestTemplateCreator
from app.services.working_ezest_template import WorkingEZestTemplateCreator

class TemplateEngine:
    """Template management and document generation engine"""
    
    def __init__(self):
        self.templates_dir = Path(settings.TEMPLATES_DIR)
        self.output_dir = Path(settings.OUTPUT_DIR)
        self.last_warnings: List[str] = []
        
        # Ensure directories exist
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        
        # Do NOT modify user's templates automatically. Only create if missing.
        main_tmpl = self.templates_dir / "ezest-updated.docx"
        if not main_tmpl.exists():
            self._ensure_ezest_updated_bullets_template()
    
    def _create_default_template(self):
        """Create a default template if none exists"""
        default_template_path = self.templates_dir / "default.docx"
        
        if not default_template_path.exists():
            # Create a basic template document
            doc = Document()
            
            # Add template placeholders
            doc.add_heading('{{contact_info.name}}', 0)
            
            # Contact information
            contact_para = doc.add_paragraph()
            contact_para.add_run('Email: {{contact_info.email}} | ')
            contact_para.add_run('Phone: {{contact_info.phone}}')
            if '{{contact_info.address}}':
                contact_para.add_run(' | {{contact_info.address}}')
            
            # Professional Summary
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph('{{summary}}')
            
            # Experience Section
            doc.add_heading('Professional Experience', level=1)
            doc.add_paragraph('''
{%- for exp in experience %}
{{ exp.title }} - {{ exp.company }}
{{ exp.start_date }} - {{ exp.end_date }}
{{ exp.description }}

{%- endfor %}
            '''.strip())
            
            # Education Section
            doc.add_heading('Education', level=1)
            doc.add_paragraph('''
{%- for edu in education %}
{{ edu.degree }}
{{ edu.institution }}
{{ edu.graduation_date }}

{%- endfor %}
            '''.strip())
            
            # Skills Section
            doc.add_heading('Skills', level=1)
            doc.add_paragraph('{{ skills | join(", ") }}')
            
            doc.save(default_template_path)
    
    def _create_ezest_template(self):
        """Create improved e-Zest template if it doesn't exist"""
        ezest_template_path = self.templates_dir / "ezest.docx"
        
        if not ezest_template_path.exists():
            try:
                # Use working template creator
                creator = WorkingEZestTemplateCreator()
                creator.create_template(ezest_template_path)
                logging.info("Working e-Zest template created successfully")
            except Exception as e:
                logging.error(f"Failed to create working e-Zest template: {str(e)}")
                # Fallback to original creator
                try:
                    creator = EZestTemplateCreator(self.templates_dir)
                    creator.create_ezest_template()
                    logging.info("Fallback e-Zest template created successfully")
                except Exception as fallback_e:
                    logging.error(f"Failed to create fallback e-Zest template: {str(fallback_e)}")
    
    def get_template_info(self, template_id: str) -> Optional[TemplateInfo]:
        """Get template information"""
        if template_id == "default":
            return TemplateInfo(
                id="default",
                name="Default Agency Template",
                description="Standard professional resume format",
                version="1.0",
                fields=["contact_info", "summary", "experience", "education", "skills"]
            )
        elif template_id == "ezest":
            return TemplateInfo(
                id="ezest",
                name="e-Zest Professional Template",
                description="Professional e-Zest formatted resume with proper fonts and spacing",
                version="1.0",
                fields=["contact_info", "summary", "experience", "education", "skills"]
            )
        elif template_id == "ezest-updated":
            return TemplateInfo(
                id="ezest-updated",
                name="e-Zest Updated Template",
                description="Updated e-Zest template with bulletized summary",
                version="1.0",
                fields=["contact_info", "summary", "summary_bullets", "experience", "education", "skills"]
            )
        elif template_id == "ezest-updated-bullets":
            return TemplateInfo(
                id="ezest-updated-bullets",
                name="e-Zest Updated Template (Bullets)",
                description="e-Zest template variant rendering summary_bullets in a loop",
                version="1.0",
                fields=["contact_info", "summary", "summary_bullets", "experience", "education", "skills"]
            )
        elif template_id == "ezest-coded":
            return TemplateInfo(
                id="ezest-coded",
                name="e-Zest Coded Template",
                description="Template generated entirely via Python (tables and formatting coded)",
                version="1.0",
                fields=[
                    "contact_info",
                    "summary",
                    "summary_bullets",
                    "experience",
                    "education",
                    "skills",
                    "other_projects",
                    "certifications_rows"
                ]
            )
        
        # Check for custom templates
        template_path = self.templates_dir / f"{template_id}.docx"
        if template_path.exists():
            return TemplateInfo(
                id=template_id,
                name=f"Custom Template {template_id}",
                description="Custom agency template",
                version="1.0",
                fields=["contact_info", "summary", "experience", "education", "skills"]
            )
        
        return None
    
    def list_templates(self) -> List[TemplateInfo]:
        """List available templates; ensure primary ones exist but never overwrite user edits."""
        # Ensure the updated bullets template exists
        main_path = self.templates_dir / "ezest-updated.docx"
        if not main_path.exists():
            try:
                self._ensure_ezest_updated_bullets_template()
            except Exception:
                pass

        # Ensure the code-generated template exists (created on demand)
        coded_path = self.templates_dir / "ezest-coded.docx"
        if not coded_path.exists():
            try:
                self._ensure_ezest_coded_template()
            except Exception:
                # Non-fatal; simply don't list it if creation failed
                pass

        templates: List[TemplateInfo] = []
        if main_path.exists():
            main = self.get_template_info("ezest-updated")
            if main:
                templates.append(main)
        if coded_path.exists():
            coded = self.get_template_info("ezest-coded")
            if coded:
                templates.append(coded)
        return templates

    def _ensure_ezest_coded_template(self) -> None:
        """Create or refresh the code-generated ezest-coded.docx using
        templates/ezest-code-gen/ezest_cv_generator.py if present.

        Behavior:
        - If target doesn't exist, generate it.
        - If target exists but generator file is newer, regenerate (safe overwrite).
        - Otherwise, no-op.
        """
        target = self.templates_dir / "ezest-coded.docx"
        # Path of the user-editable generator
        gen_path = (Path(__file__).resolve().parents[2] / "templates" / "ezest-code-gen" / "ezest_cv_generator.py")
        if not gen_path.exists():
            logging.warning("Generator not found at templates/ezest-code-gen/ezest_cv_generator.py; skipping coded template creation")
            return

        # If target exists and is newer than generator, skip regeneration
        try:
            if target.exists():
                if target.stat().st_mtime >= gen_path.stat().st_mtime:
                    return
        except Exception:
            # If stat fails for any reason, proceed to attempt regeneration
            pass
        try:
            spec = importlib.util.spec_from_loader("ezest_cv_generator", SourceFileLoader("ezest_cv_generator", str(gen_path)))
            if spec is None or spec.loader is None:
                logging.warning("Could not load ezest_cv_generator spec")
                return
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)  # type: ignore
            if hasattr(module, "EZestCVTemplateGenerator"):
                Generator = getattr(module, "EZestCVTemplateGenerator")
                generator = Generator()
                generator.create_complete_template(target)
                logging.info("(Re)created code-generated template ezest-coded.docx from templates/ezest-code-gen/ezest_cv_generator.py")
            else:
                logging.warning("EZestCVTemplateGenerator not found in ezest_cv_generator module")
        except Exception as e:
            logging.error(f"Failed to create ezest-coded template: {e}")
    
    def apply_template(self, extracted_data: ExtractedData, template_id: str) -> str:
        """Apply template to extracted data and generate formatted document"""
        try:
            self.last_warnings = []
            template_path = self.templates_dir / f"{template_id}.docx"

            if not template_path.exists():
                raise FileNotFoundError(f"Template not found: {template_id}")

            # Prepare context data first (used both for population and rendering)
            context = self._prepare_template_context(extracted_data)
            # Simple required content checks and warnings
            self._collect_warnings(context)

            # If this is the code-generated template, pre-populate rows before rendering
            if template_id == "ezest-coded":
                # Copy template to a temp path for in-place population prior to docxtpl rendering
                temp_filename = f"_tmp_{uuid.uuid4().hex[:8]}_{template_id}.docx"
                temp_path = self.output_dir / temp_filename
                shutil.copyfile(template_path, temp_path)

                # Load the temp document and populate tables based on headers
                doc = Document(temp_path)
                try:
                    gen_path = (Path(__file__).resolve().parents[2] / "templates" / "ezest-code-gen" / "ezest_cv_generator.py")
                    spec = importlib.util.spec_from_loader("ezest_cv_generator", SourceFileLoader("ezest_cv_generator", str(gen_path)))
                    if spec and spec.loader:
                        module = importlib.util.module_from_spec(spec)
                        spec.loader.exec_module(module)  # type: ignore
                        if hasattr(module, "EZestCVTemplateGenerator"):
                            Generator = getattr(module, "EZestCVTemplateGenerator")
                            generator = Generator()
                            # Populate sections by header markers
                            generator.populate_skills(doc, context.get('skills_rows', []))
                            generator.populate_other_projects(doc, context.get('other_projects', []))
                            generator.populate_education(doc, context.get('education', []))
                            generator.populate_certifications(doc, context.get('certifications_rows', []))
                        else:
                            logging.warning("EZestCVTemplateGenerator not found; skipping row population")
                    else:
                        logging.warning("Could not load ezest_cv_generator spec; skipping row population")
                except Exception as e:
                    logging.warning(f"Row population step failed, proceeding with plain rendering: {e}")
                # Save populated temp doc
                doc.save(temp_path)

                # Now render with docxtpl using the populated temp doc
                template = DocxTemplate(temp_path)
            else:
                # Non-coded templates: render directly
                template = DocxTemplate(template_path)

            # Render template
            template.render(context)

            # Generate output filename
            output_filename = f"formatted_{uuid.uuid4().hex[:8]}_{template_id}.docx"
            output_path = self.output_dir / output_filename

            # Save rendered document
            template.save(output_path)

            return output_filename

        except Exception as e:
            logging.error(f"Template application failed: {str(e)}")
            raise ValueError(f"Failed to apply template: {str(e)}")
    
    def _prepare_template_context(self, extracted_data: ExtractedData) -> Dict[str, Any]:
        """Prepare context data for template rendering"""
        # Bulletize summary into list items suitable for templating
        # Prefer LLM-provided bullets if available in additional_data
        addl = getattr(extracted_data, 'additional_data', {}) or {}
        llm_bullets = addl.get('summary_bullets') if isinstance(addl, dict) else None
        if llm_bullets and isinstance(llm_bullets, list) and any(str(b).strip() for b in llm_bullets):
            bullets = [str(b).strip() for b in llm_bullets if str(b).strip()]
            # Also synthesize a plain summary text for any template using a text block
            synthesized_summary = '; '.join(bullets)
        else:
            bullets = self._bulletize_summary(extracted_data.summary or '')
            synthesized_summary = extracted_data.summary or ''
        # Build Tools & Technologies helpers
        skills_grouped = getattr(extracted_data, 'skills_grouped', {}) or {}
        tools_title = getattr(extracted_data, 'tools_title', None) or 'Professional Skills'
        skills_left_lines = ''
        skills_right_lines = ''
        skills_rows: List[Dict[str, str]] = []
        if skills_grouped:
            left_lines: List[str] = []
            right_lines: List[str] = []
            for label, items in skills_grouped.items():
                left_lines.append(str(label))
                right_lines.append(', '.join([str(x) for x in (items or []) if str(x).strip()]))
                skills_rows.append({
                    'left': str(label),
                    'right': ', '.join([str(x) for x in (items or []) if str(x).strip()])
                })
            skills_left_lines = '\n'.join(left_lines)
            skills_right_lines = '\n'.join(right_lines)
        else:
            # Fallback: keep single row using flat skills
            skills_left_lines = 'Skills'
            skills_right_lines = ', '.join(extracted_data.skills or [])
            if skills_right_lines:
                skills_rows.append({'left': skills_left_lines, 'right': skills_right_lines})

        # Title: ensure empty string if missing so no stray characters render
        safe_title = (getattr(extracted_data.contact_info, 'title', None) or '').strip()

        # Build simple rows for Other Notable Projects / Education / Certifications (for templates expecting simple loops)
        simple_other_projects: List[Dict[str, Any]] = []
        for p in addl.get('other_notable_projects', []) or []:
            try:
                # Normalize technologies to list of strings; join in template
                techs = p.get('technologies') or p.get('technology') or []
                if isinstance(techs, str):
                    techs = [techs]
                simple_other_projects.append({
                    'name': p.get('project_name') or p.get('name') or '',
                    'duration': p.get('duration') or '',
                    'technologies': [str(x).strip() for x in techs if str(x).strip()],
                })
            except Exception:
                continue

        simple_education: List[Dict[str, Any]] = []
        for e in addl.get('education', []) or []:
            try:
                simple_education.append({
                    'degree': e.get('degree') or e.get('course') or '',
                    'institution': e.get('institution') or e.get('university') or e.get('board') or '',
                    'graduation_date': e.get('year') or e.get('graduation_date') or '',
                })
            except Exception:
                continue

        simple_cert_rows: List[Dict[str, Any]] = []
        c_list = addl.get('certifications', []) or []
        for i, c in enumerate(c_list, start=1):
            try:
                authority = c.get('issuer') or c.get('authority') or c.get('name') or ''
                simple_cert_rows.append({'sno': i, 'authority': authority})
            except Exception:
                continue

        context = {
            'contact_info': {
                'name': extracted_data.contact_info.name or 'N/A',
                'email': extracted_data.contact_info.email or 'N/A',
                'phone': extracted_data.contact_info.phone or 'N/A',
                'address': extracted_data.contact_info.address or '',
                'linkedin': extracted_data.contact_info.linkedin or '',
                'website': extracted_data.contact_info.website or '',
                'title': safe_title
            },
            # Keep previous behavior: prefer bullet list; also provide a reasonable text fallback
            'summary': synthesized_summary or 'Professional summary not available.',
            'summary_bullets': bullets,
            'experience': [],
            # Backward-compatible education loop and new alias
            'education': simple_education,
            'skills': extracted_data.skills or [],
            'tools_title': tools_title,
            'skills_grouped': skills_grouped,
            # Preformatted 2-column lines for the template (no style changes needed)
            'skills_left_lines': skills_left_lines,
            'skills_right_lines': skills_right_lines,
            # Row-wise pairs for docxtpl row loop
            'skills_rows': skills_rows,
            # Enhanced variables exposed for templates that use the new structure
            'detailed_experience': addl.get('detailed_experience', []),
            'other_notable_projects': addl.get('other_notable_projects', []),
            'education_table': addl.get('education', []),
            'certifications': addl.get('certifications', []),
            # Simple aliases to match the user's current template
            'other_projects': simple_other_projects,
            'certifications_rows': simple_cert_rows,
        }
        
        # Process experience data
        # If enhanced detailed_experience is available, map it to the legacy 'experience' loop format
        enhanced_exps = addl.get('detailed_experience') if isinstance(addl, dict) else None
        if isinstance(enhanced_exps, list) and enhanced_exps:
            for e in enhanced_exps:
                try:
                    duration = str(e.get('duration') or '').strip()
                    start_date = ''
                    end_date = ''
                    if ' - ' in duration:
                        start_date, end_date = [s.strip() for s in duration.split(' - ', 1)]
                    elif duration:
                        start_date = duration
                    techs = e.get('technologies_used') or e.get('technology') or []
                    if isinstance(techs, str):
                        techs = [techs]
                    achievements = e.get('key_achievements') or []
                    if isinstance(achievements, str):
                        achievements = [achievements]
                    # Prefer explicit project description field; otherwise synthesize from first achievement
                    proj_desc = e.get('project_description') or e.get('description') or ''
                    if not proj_desc and achievements:
                        proj_desc = achievements[0]
                    context['experience'].append({
                        'title': e.get('project_name') or e.get('title') or 'Project',
                        'company': e.get('organization') or e.get('company') or '',
                        'location': e.get('location') or '',
                        'start_date': start_date or '—',
                        'end_date': end_date or 'Present',
                        'description': proj_desc or '',
                        'project_description': proj_desc or '',
                        'technologies': [str(x) for x in techs if str(x).strip()],
                        'responsibilities': [str(x).strip() for x in achievements if str(x).strip()],
                        'is_current': True if (end_date.lower() == 'present') else False
                    })
                except Exception:
                    continue
            # Build overflow -> other_projects (items beyond top 5)
            try:
                overflow = enhanced_exps[5:] if len(enhanced_exps) > 5 else []
                extra_rows = []
                for p in overflow:
                    dur = p.get('duration') or ''
                    techs = p.get('technologies_used') or p.get('technology') or []
                    if isinstance(techs, str):
                        techs = [techs]
                    extra_rows.append({
                        'name': p.get('project_name') or p.get('title') or '',
                        'duration': dur,
                        'technologies': [str(x).strip() for x in techs if str(x).strip()]
                    })
                # Merge with any existing simple rows derived earlier or from LLM other_notable_projects
                existing_rows = context.get('other_projects') or []
                context['other_projects'] = (existing_rows + extra_rows) if extra_rows else existing_rows
            except Exception:
                pass
            # After mapping enhanced format, skip legacy parsing below
            return context

        # Build a lightweight vocabulary from provided skills/groups
        skills_vocab: set = set()
        try:
            for k, v in (skills_grouped or {}).items():
                for s in (v or []):
                    if isinstance(s, str) and s.strip():
                        skills_vocab.add(s.strip().lower())
            for s in (extracted_data.skills or []):
                if isinstance(s, str) and s.strip():
                    skills_vocab.add(s.strip().lower())
        except Exception:
            skills_vocab = set()

        def first_sentence(text: str, max_len: int = 300) -> str:
            t = (text or '').strip()
            if not t:
                return ''
            # Split on sentence enders; fallback to truncation
            import re
            parts = re.split(r'(?<=[.!?])\s+', t)
            s = parts[0] if parts else t
            if len(s) > max_len:
                s = s[:max_len].rstrip() + '...'
            return s

        def project_level_summary(text: str) -> str:
            """Try to pick a project-level sentence (what the project is), not actions."""
            import re
            t = (text or '').strip().lstrip('•- ').strip()
            if not t:
                return ''
            sentences = re.split(r'(?<=[.!?])\s+', t)
            action_verbs = {
                'led','designed','implemented','configured','developed','built','created','maintained',
                'managed','troubleshot','integrated','optimized','handled','validated','deployed','architected'
            }
            for s in sentences:
                s_clean = s.strip()
                if not s_clean:
                    continue
                first_word = re.split(r'\W+', s_clean.lower())[0]
                if first_word not in action_verbs and len(s_clean) > 20:
                    return s_clean
            # Fallback to first sentence without the bullet
            return first_sentence(t)

        def extract_technologies(text: str) -> List[str]:
            if not text:
                return []
            import re
            # Known tech/tool whitelist keywords (not exhaustive, can be expanded)
            TECH_WHITELIST = {
                'aws','azure','gcp','kubernetes','docker','terraform','ansible','jenkins','gitlab','github',
                'python','java','c#','.net','node','react','angular','sql','postgres','mysql','mssql','oracle',
                'snowflake','spark','hadoop','airflow','kafka','rabbitmq','redis','mongo','elasticsearch','kibana',
                'grafana','prometheus','tableau','powerbi','pandas','numpy','scikit','tensorflow','pytorch',
                'rest','soap','graphql','grpc','s3','ec2','rds','eks','aks','gke',
                'servicenow','workday','scom','sccm','scorch','scsm','okta','splunk','sonarqube','vault',
                'powershell','bash','linux','windows','nginx','apache'
            }
            STOP = {
                'the','and','for','with','using','to','of','in','on','by','a','an','at','from','via','including',
                'system','systems','solution','solutions','service','services','tools','technology','technologies',
                'center','desk','team','teams','process','processes','module','modules','api','apis','reports'
            }
            # Collect candidates
            tokens_all = re.findall(r'[A-Za-z0-9+.#-]+', text)
            candidates = []
            for w in tokens_all:
                lw = w.lower()
                if lw in STOP:
                    continue
                if lw in TECH_WHITELIST or lw in skills_vocab:
                    candidates.append(lw)
                # also include clear acronyms (>=2 uppercase letters)
                elif w.isupper() and len(w) >= 2:
                    candidates.append(lw)
            # Map back to nicely-cased labels using skills vocab originals if possible
            techs: List[str] = []
            uniq = []
            for t in candidates:
                # Try to find original-cased skill
                match = next((sv for sv in skills_vocab if sv == t), None)
                label = match or t
                if label not in uniq:
                    uniq.append(label)
            return uniq

        def bulletize_responsibilities(text: str) -> List[str]:
            if not text:
                return []
            import re
            # Split on bullet markers or newlines
            raw = re.split(r'\n|•|-\s+', text)
            items: List[str] = []
            for it in raw:
                s = it.strip(' •\t-\r')
                if not s:
                    continue
                # Drop lines that are mostly pure tech lists (3+ commas)
                if s.count(',') >= 3:
                    # keep as tech list, not responsibility
                    continue
                items.append(s)
            return items

        for exp in extracted_data.experience:
            desc = (exp.description or '').strip()
            techs = extract_technologies(desc)
            resp = bulletize_responsibilities(desc)
            project_desc = project_level_summary(desc)

            experience_item = {
                'title': exp.title or 'Position Title',
                'company': exp.company or 'Company Name',
                'location': exp.location or '',
                'start_date': exp.start_date or 'Start Date',
                'end_date': exp.end_date or 'End Date',
                # Keep original full description for reference/fallback
                'description': desc or 'Job description not available.',
                # New fields for template
                'project_description': project_desc or (desc[:200] + '...') if desc else '—',
                'technologies': techs,
                'responsibilities': resp if resp else [],
                'is_current': exp.is_current
            }
            context['experience'].append(experience_item)
        
        # Process education data
        for edu in extracted_data.education:
            education_item = {
                'degree': edu.degree or 'Degree',
                'institution': edu.institution or 'Institution',
                'location': edu.location or '',
                'graduation_date': edu.graduation_date or 'Graduation Date',
                'gpa': edu.gpa or '',
                'honors': edu.honors or ''
            }
            context['education'].append(education_item)
        
        return context
    
    def validate_template(self, template_path: Path) -> Dict[str, Any]:
        """Validate template structure and required fields"""
        try:
            template = DocxTemplate(template_path)
            
            # Prepare a comprehensive sample context to validate placeholders
            sample_context: Dict[str, Any] = {
                'contact_info': {
                    'name': 'John Doe',
                    'email': 'john.doe@example.com',
                    'phone': '+1-555-123-4567',
                    'address': '123 Main St, City',
                    'linkedin': 'linkedin.com/in/johndoe',
                    'website': 'johndoe.dev',
                },
                'summary': 'Seasoned professional with expertise in ...',
                'experience': [
                    {
                        'title': 'Senior Developer',
                        'company': 'Acme Corp',
                        'location': 'Remote',
                        'start_date': 'Jan 2020',
                        'end_date': 'Present',
                        'description': 'Built things',
                        'is_current': True,
                    }
                ],
                'education': [
                    {
                        'degree': 'B.Sc. Computer Science',
                        'institution': 'Tech University',
                        'location': 'City',
                        'graduation_date': '2018',
                        'gpa': '3.9',
                        'honors': 'Summa Cum Laude',
                    }
                ],
                'skills': ['Python', 'FastAPI', 'React']
            }

            validation_result: Dict[str, Any] = {
                'valid': True,
                'errors': [],
                'warnings': [],
                'required_fields': [
                    'contact_info.name', 'contact_info.email', 'contact_info.phone',
                    'experience', 'education', 'skills'
                ],
                'found_fields': [],
                'undeclared_variables': []
            }

            # Compute undeclared variables using docxtpl to catch typos/mistakes
            try:
                undeclared = template.get_undeclared_template_variables(sample_context)
            except Exception as e:
                undeclared = set()
                validation_result['warnings'].append(f"Could not compute undeclared variables: {str(e)}")

            if undeclared:
                validation_result['undeclared_variables'] = sorted(list(undeclared))
                validation_result['valid'] = False
                validation_result['errors'].append(
                    "Template references variables not present in the supported context"
                )

            # Read all visible text (paragraphs + tables) to detect presence of required placeholders
            doc = Document(template_path)
            template_text_parts: List[str] = []
            for paragraph in doc.paragraphs:
                template_text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            template_text_parts.append(para.text)
            template_text = "\n".join(template_text_parts)

            for var in validation_result['required_fields']:
                if var in template_text:
                    validation_result['found_fields'].append(var)

            # Add warnings for missing commonly expected fields
            missing_fields = set(validation_result['required_fields']) - set(validation_result['found_fields'])
            for field in sorted(list(missing_fields)):
                validation_result['warnings'].append(
                    f"Template variable '{{{{{field}}}}}' not found in paragraphs/tables"
                )

            return validation_result
        
        except Exception as e:
            return {
                'valid': False,
                'errors': [f"Template validation failed: {str(e)}"],
                'warnings': [],
                'required_fields': [],
                'found_fields': [],
                'undeclared_variables': []
            }
    
    def create_template_from_sample(self, sample_docx_path: Path, template_id: str) -> bool:
        """Create a new template from a sample document"""
        try:
            # Copy sample to templates directory
            new_template_path = self.templates_dir / f"{template_id}.docx"
            
            # Read the sample document
            doc = Document(sample_docx_path)
            
            # This is a simplified version - in practice, you'd want to:
            # 1. Analyze the document structure
            # 2. Identify sections that should be templated
            # 3. Replace content with appropriate template variables
            
            # For now, just copy the document and add basic template variables
            doc.save(new_template_path)
            
            return True
            
        except Exception as e:
            logging.error(f"Template creation failed: {str(e)}")
            return False
    
    def get_template_preview(self, template_id: str) -> Dict[str, Any]:
        """Get template structure for preview"""
        template_info = self.get_template_info(template_id)
        if not template_info:
            return {}
        
        return {
            'template_id': template_id,
            'name': template_info.name,
            'structure': {
                'header': {
                    'name': '{{contact_info.name}}',
                    'contact': '{{contact_info.email}} | {{contact_info.phone}}'
                },
                'sections': [
                    {
                        'title': 'Professional Summary',
                        'content': '{% if summary_bullets and summary_bullets|length > 0 %}{% for s in summary_bullets %}• {{ s }}{% endfor %}{% else %}{{summary}}{% endif %}'
                    },
                    {
                        'title': 'Professional Experience',
                        'content': '{% for exp in experience %}{{exp.title}} - {{exp.company}}{% endfor %}'
                    },
                    {
                        'title': 'Education',
                        'content': '{% for edu in education %}{{edu.degree}} - {{edu.institution}}{% endfor %}'
                    },
                    {
                        'title': 'Skills',
                        'content': '{{skills | join(", ")}}'
                    }
                ]
            }
        }
    
    def _bulletize_summary(self, summary: str) -> List[str]:
        """Convert summary text into bullet points. Prefer AI formatting if configured, otherwise heuristic split."""
        cleaned = (summary or '').strip()
        if not cleaned:
            return []
        # Heuristic split: prefer newlines, then sentence boundaries
        parts: List[str] = []
        # Normalize whitespace
        cleaned = cleaned.replace('\r', '')
        # Split by explicit newlines first
        for line in cleaned.split('\n'):
            line = line.strip(' •-\t')
            if not line:
                continue
            # Further split long lines by sentences
            sentences = re.split(r'(?<=[.!?])\s+', line)
            for s in sentences:
                s = s.strip(' •-\t')
                if len(s) >= 2:
                    parts.append(s)
        # Deduplicate and cap length
        uniq = []
        seen = set()
        for p in parts:
            k = p.lower()
            if k not in seen:
                uniq.append(p)
                seen.add(k)
        return uniq[:10]
    
    def _collect_warnings(self, context: Dict[str, Any]) -> None:
        """Populate self.last_warnings based on missing or weak content for templates."""
        warnings: List[str] = []
        ci = context.get('contact_info', {}) or {}
        if not ci.get('name') or ci.get('name') == 'N/A':
            warnings.append('Missing candidate name')
        if not ci.get('email') or ci.get('email') == 'N/A':
            warnings.append('Missing email')
        if not context.get('summary') or 'not available' in str(context.get('summary')).lower():
            warnings.append('Missing professional summary')
        if not context.get('summary_bullets'):
            warnings.append('Summary not bulletized or missing')
        if not context.get('skills'):
            warnings.append('Skills section is empty')
        # Keep for retrieval after render
        self.last_warnings = warnings
    
    def get_last_warnings(self) -> List[str]:
        return self.last_warnings
    
    def _prune_templates_to_main(self) -> None:
        """Move all templates except the main ezest-updated.docx to templates/backup for a clean single-template setup."""
        backup_dir = self.templates_dir / "backup"
        backup_dir.mkdir(exist_ok=True)
        try:
            for docx in self.templates_dir.glob("*.docx"):
                # Keep only the main template in place
                if docx.name.lower() == "ezest-updated.docx":
                    continue
                # Backup others with a unique name if needed
                dest = backup_dir / docx.name
                if dest.exists():
                    ts = datetime.now().strftime('%Y%m%d%H%M%S')
                    dest = backup_dir / f"{docx.stem}_{ts}{docx.suffix}"
                try:
                    docx.replace(dest)
                except Exception as be:
                    logging.warning(f"Could not backup {docx.name}: {be}")
        except Exception as e:
            logging.warning(f"Prune-to-main encountered an issue: {e}")

    def _ensure_ezest_updated_bullets_template(self) -> None:
        """Create ezest-updated.docx only if it does not exist. Never overwrite user-updated template."""
        try:
            target = self.templates_dir / "ezest-updated.docx"
            if target.exists():
                logging.info("ezest-updated.docx already exists; not recreating or modifying it")
                return
            # Build fresh template
            doc = Document()
            # Header: Candidate Name and Title
            heading = doc.add_paragraph()
            heading.add_run('{{contact_info.name}}').bold = True
            title_p = doc.add_paragraph()
            title_p.add_run('({{ contact_info.title | default("Professional Title") }})')
            
            # Profile Summary (bullets)
            doc.add_heading('Profile Summary', level=1)
            doc.add_paragraph('{% if summary_bullets and summary_bullets|length > 0 %}')
            doc.add_paragraph('{% for single_line in summary_bullets %}• {{ single_line }}{% endfor %}')
            doc.add_paragraph('{% else %}{{ summary }}{% endif %}')
            
            # Professional Skills (2-column table)
            doc.add_heading('Professional Skills', level=1)
            skills_tbl = doc.add_table(rows=1, cols=2)
            skills_tbl.autofit = True
            skills_tbl.cell(0,0).text = '{% for skill in skills[0::2] %}{{ skill }}{% if not loop.last %}, {% endif %}{% endfor %}'
            skills_tbl.cell(0,1).text = '{% for skill in skills[1::2] %}{{ skill }}{% if not loop.last %}, {% endif %}{% endfor %}'
            
            # Relevant Work Experience (single-column table with bullets for responsibilities)
            doc.add_heading('Relevant Work Experience', level=1)
            exp_tbl = doc.add_table(rows=1, cols=1)
            exp_tbl.cell(0,0).text = (
                '{% for exp in experience %}'
                'Project #{{ loop.index }}\n'
                'Duration: {{ exp.start_date }} - {{ exp.end_date }}\n\n'
                'Project Description: {{ exp.description }}\n\n'
                'Technology: {{ exp.technologies | join(", ") if exp.technologies else "" }}\n\n'
                'Role & Responsibilities:\n'
                '{% if exp.responsibilities %}'
                '{% for r in exp.responsibilities %}• {{ r }}\n{% endfor %}'
                '{% else %}• {{ exp.description }}\n{% endif %}\n\n'
                '{% endfor %}'
            )
            
            # Other Notable Projects (optional 3-column table)
            doc.add_heading('Other Notable Projects', level=1)
            onp = doc.add_table(rows=2, cols=3)
            onp.cell(0,0).text = 'Project'
            onp.cell(0,1).text = 'Duration'
            onp.cell(0,2).text = 'Technology'
            onp.cell(1,0).text = '{% if other_projects %}{% for p in other_projects %}{{ p.name }}{% if not loop.last %}\n{% endif %}{% endfor %}{% else %}{% endif %}'
            onp.cell(1,1).text = '{% if other_projects %}{% for p in other_projects %}{{ p.duration }}{% if not loop.last %}\n{% endif %}{% endfor %}{% else %}{% endif %}'
            onp.cell(1,2).text = '{% if other_projects %}{% for p in other_projects %}{{ p.technologies | join(", ") }}{% if not loop.last %}\n{% endif %}{% endfor %}{% else %}{% endif %}'
            
            # Education Details (3 columns)
            doc.add_heading('Education Details', level=1)
            edu_tbl = doc.add_table(rows=2, cols=3)
            edu_tbl.cell(0,0).text = 'Course'
            edu_tbl.cell(0,1).text = 'University'
            edu_tbl.cell(0,2).text = 'Year of Passing'
            edu_tbl.cell(1,0).text = '{% for edu in education %}{{ edu.degree }}{% if not loop.last %}\n{% endif %}{% endfor %}'
            edu_tbl.cell(1,1).text = '{% for edu in education %}{{ edu.institution }}{% if not loop.last %}\n{% endif %}{% endfor %}'
            edu_tbl.cell(1,2).text = '{% for edu in education %}{{ edu.graduation_date }}{% if not loop.last %}\n{% endif %}{% endfor %}'
            
            # Save default template (only when missing)
            doc.save(target)
            logging.info("Created ezest-updated.docx because it was missing")
        except Exception as e:
            logging.error(f"Failed to (re)create ezest-updated template: {e}")
