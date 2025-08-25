from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
import logging

class ImprovedEZestTemplateCreator:
    """Creates an improved e-Zest template matching the HTML structure"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def create_template(self, template_path: Path):
        """Create improved e-Zest template with proper table structure"""
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # Header - Name and Title
            self._add_header(doc)
            
            # Profile Summary Section
            self._add_profile_summary(doc)
            
            # Professional Skills Section
            self._add_professional_skills(doc)
            
            # Relevant Work Experience Section
            self._add_work_experience(doc)
            
            # Other Notable Projects Section
            self._add_notable_projects(doc)
            
            # Education Details Section
            self._add_education_details(doc)
            
            # Save template
            doc.save(template_path)
            self.logger.info(f"Improved e-Zest template created at {template_path}")
            
        except Exception as e:
            self.logger.error(f"Error creating improved e-Zest template: {str(e)}")
            raise
    
    def _add_header(self, doc):
        """Add header with name and title"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('{{contact_info.name}}')
        name_run.font.size = Pt(16)
        name_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)  # #17365D
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Image
        try:
            doc.add_picture('/mnt/d/Projects/KP HR Resume Formatter/docs/Sample files/CV_Name_Title_e-Zest 1 (1),image001/CV_Name_Title_e-Zest 1 (1)_files/image001.png', width=Inches(6.5))
        except FileNotFoundError:
            self.logger.warning("Header image not found, skipping.")
            pass
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('({{contact_info.title | default("Professional Title")}})')
        title_run.font.size = Pt(10)
        title_run.font.name = 'Segoe UI'
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add some spacing
        doc.add_paragraph()
    
    def _add_profile_summary(self, doc):
        """Add Profile Summary section with bulleted table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Profile Summary')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create table for summary
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table borders
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        cell = table.cell(0, 0)
        cell.text = '{{summary}}'
        
        doc.add_paragraph()
    
    def _add_professional_skills(self, doc):
        """Add Professional Skills section with 2-column table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Professional Skills')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create 2-column table for skills
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.2)
        
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        # Add a template loop for skills
        table.cell(0, 0).text = "{% for skill in skills %}{{ skill.category }}{% endfor %}"
        table.cell(0, 1).text = "{% for skill in skills %}{{ skill.skills|join(', ') }}{% endfor %}"
        
        doc.add_paragraph()
    
    def _add_work_experience(self, doc):
        """Add Relevant Work Experience section"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Relevant Work Experience')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create table for each experience entry
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        cell = table.cell(0, 0)
        cell.text = """
{% for exp in experience %}
Project #{{ loop.index }}                                                                                             Duration: {{ exp.start_date }} - {{ exp.end_date }}
Project Description: {{ exp.description }}
Technology: {{ exp.technologies|join(', ') }}
Role & Responsibilities: 
{% for resp in exp.responsibilities %}o {{ resp }}
{% endfor %}
{% endfor %}
"""
        doc.add_paragraph()
    
    def _add_notable_projects(self, doc):
        """Add Other Notable Projects section with 3-column table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Other Notable Projects')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create 3-column table
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        # Header row with blue background
        header_cells = ['Project', 'Duration', 'Technology']
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            self._set_cell_background(cell, RGBColor(0x17, 0x36, 0x5D))
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        table.add_row().cells[0].text = "{% for project in other_projects %}{{ project.name }}{% endfor %}"
        table.add_row().cells[1].text = "{% for project in other_projects %}{{ project.duration }}{% endfor %}"
        table.add_row().cells[2].text = "{% for project in other_projects %}{{ project.technologies|join(', ') }}{% endfor %}"
        
        doc.add_paragraph()
    
    def _add_education_details(self, doc):
        """Add Education Details section with 3-column table"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('Education Details')
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        heading_run.bold = True
        
        # Create 3-column table
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table, RGBColor(0xBF, 0xBF, 0xBF))
        
        # Header row with blue background
        header_cells = ['Course', 'University', 'Year of Passing']
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            self._set_cell_background(cell, RGBColor(0x17, 0x36, 0x5D))
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        table.add_row().cells[0].text = "{% for edu in education %}{{ edu.degree }}{% endfor %}"
        table.add_row().cells[1].text = "{% for edu in education %}{{ edu.institution }}{% endfor %}"
        table.add_row().cells[2].text = "{% for edu in education %}{{ edu.graduation_year }}{% endfor %}"
        
        doc.add_paragraph()
    
    def _set_table_borders(self, table, color):
        """Set table borders with specified color"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = tcBorders.find(qn(f"w:{border_name}"))
                    if border is None:
                        border = OxmlElement(f'w:{border_name}')
                        tcBorders.append(border)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:color'), str(color))
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:fill'), str(color))