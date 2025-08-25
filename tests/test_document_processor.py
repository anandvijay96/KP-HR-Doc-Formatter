import pytest
import asyncio
from pathlib import Path
from unittest.mock import Mock, patch, AsyncMock
from docx import Document
import tempfile
import os

from app.services.document_processor import DocumentProcessor
from app.models.schemas import ExtractedData, ContactInfo

class TestDocumentProcessor:
    """Test cases for DocumentProcessor"""
    
    @pytest.fixture
    def processor(self):
        return DocumentProcessor()
    
    @pytest.fixture
    def sample_docx_content(self):
        """Create a sample DOCX file for testing"""
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("john.doe@email.com")
        doc.add_paragraph("(555) 123-4567")
        doc.add_paragraph("")
        doc.add_paragraph("EXPERIENCE")
        doc.add_paragraph("Software Engineer - Tech Corp")
        doc.add_paragraph("2020 - Present")
        doc.add_paragraph("Developed web applications using Python and React")
        doc.add_paragraph("")
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("Python, JavaScript, React, SQL")
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            return tmp.name
    
    @pytest.mark.asyncio
    async def test_save_uploaded_file(self, processor):
        """Test file upload and saving"""
        file_content = b"test content"
        filename = "test_resume.docx"
        
        unique_filename = await processor.save_uploaded_file(file_content, filename)
        
        assert unique_filename.endswith("_test_resume.docx")
        assert len(unique_filename) > len(filename)  # Should have UUID prefix
        
        # Cleanup
        file_path = processor.upload_dir / unique_filename
        if file_path.exists():
            file_path.unlink()
    
    @pytest.mark.asyncio
    async def test_save_uploaded_file_invalid_extension(self, processor):
        """Test file upload with invalid extension"""
        file_content = b"test content"
        filename = "test_resume.txt"
        
        with pytest.raises(ValueError, match="Unsupported file type"):
            await processor.save_uploaded_file(file_content, filename)
    
    def test_extract_text_from_docx(self, processor, sample_docx_content):
        """Test text extraction from DOCX file"""
        file_path = Path(sample_docx_content)
        
        text = processor.extract_text_from_docx(file_path)
        
        assert "John Doe" in text
        assert "john.doe@email.com" in text
        assert "Software Engineer" in text
        assert "Python, JavaScript" in text
        
        # Cleanup
        os.unlink(sample_docx_content)
    
    def test_basic_data_extraction(self, processor):
        """Test basic data extraction functionality"""
        sample_text = """
        John Doe
        john.doe@email.com
        (555) 123-4567
        
        EXPERIENCE
        Software Engineer - Tech Corp
        2020 - Present
        
        SKILLS
        Python, JavaScript, React
        """
        
        extracted_data = processor.basic_data_extraction(sample_text)
        
        assert isinstance(extracted_data, ExtractedData)
        assert extracted_data.contact_info.name == "John Doe"
        assert extracted_data.contact_info.email == "john.doe@email.com"
        assert extracted_data.contact_info.phone == "5551234567"
        assert len(extracted_data.skills) > 0
        assert "Python" in extracted_data.skills
    
    @patch('app.services.document_processor.NLPExtractor')
    def test_advanced_data_extraction(self, mock_nlp_extractor, processor):
        """Test advanced data extraction with NLP"""
        # Mock NLP extractor
        mock_extractor = Mock()
        mock_extractor.extract_contact_info.return_value = ContactInfo(
            name="John Doe",
            email="john.doe@email.com",
            phone="555-123-4567"
        )
        mock_extractor.extract_experience.return_value = []
        mock_extractor.extract_education.return_value = []
        mock_extractor.extract_skills.return_value = ["Python", "JavaScript"]
        
        mock_nlp_extractor.return_value = mock_extractor
        
        # Reinitialize processor to use mocked NLP extractor
        processor = DocumentProcessor()
        
        sample_text = "John Doe\njohn.doe@email.com\n555-123-4567"
        extracted_data = processor.advanced_data_extraction(sample_text)
        
        assert extracted_data.contact_info.name == "John Doe"
        assert extracted_data.contact_info.email == "john.doe@email.com"
        assert "Python" in extracted_data.skills
    
    def test_calculate_confidence_score(self, processor):
        """Test confidence score calculation"""
        contact_info = ContactInfo(
            name="John Doe",
            email="john.doe@email.com",
            phone="555-123-4567"
        )
        experience = [Mock()]  # One experience item
        education = [Mock()]   # One education item
        skills = ["Python", "JavaScript", "React", "SQL"]  # Multiple skills
        
        score = processor._calculate_confidence_score(
            contact_info, experience, education, skills
        )
        
        assert 0.0 <= score <= 1.0
        assert score > 0.5  # Should be reasonably high with good data
    
    def test_extract_summary(self, processor):
        """Test summary extraction"""
        sample_text = """
        John Doe
        
        PROFESSIONAL SUMMARY
        Experienced software engineer with 5 years of experience
        in web development and system design.
        
        EXPERIENCE
        Software Engineer - Tech Corp
        """
        
        summary = processor._extract_summary(sample_text)
        
        assert summary is not None
        assert "Experienced software engineer" in summary
        assert "5 years" in summary
    
    @pytest.mark.asyncio
    async def test_process_document_file_not_found(self, processor):
        """Test processing non-existent file"""
        with pytest.raises(FileNotFoundError):
            await processor.process_document("nonexistent_file.docx")
    
    def test_cleanup_file(self, processor):
        """Test file cleanup functionality"""
        # Create a temporary file
        test_file = processor.upload_dir / "test_cleanup.txt"
        test_file.write_text("test content")
        
        assert test_file.exists()
        
        # Test cleanup
        result = processor.cleanup_file("test_cleanup.txt")
        
        assert result is True
        assert not test_file.exists()
    
    def test_cleanup_nonexistent_file(self, processor):
        """Test cleanup of non-existent file"""
        result = processor.cleanup_file("nonexistent_file.txt")
        assert result is False
