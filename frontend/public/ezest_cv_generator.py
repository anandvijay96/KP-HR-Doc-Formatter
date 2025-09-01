from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
import logging


class EZestCVTemplateGenerator:
    """
    Complete e-Zest CV Template Generator with global styling variables
    Based on Rahul Shrivastav CV format with all sections
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # GLOBAL STYLING VARIABLES - Change these to apply across entire template
        self.HEADING_FONT = 'Calibri'          # Main section headings font
        self.BODY_FONT = 'Segoe UI'            # Body text font
        self.NAME_FONT_SIZE = Pt(16)           # Candidate name size
        self.TITLE_FONT_SIZE = Pt(12)          # Job title size
        self.HEADING_FONT_SIZE = Pt(14)        # Section headings size
        self.BODY_FONT_SIZE = Pt(10)           # Regular body text size
        self.TABLE_HEADER_FONT_SIZE = Pt(11)   # Table headers size
        
        # COLORS
        self.HEADING_COLOR = RGBColor(0x17, 0x36, 0x5D)      # Dark blue for headings
        self.BODY_COLOR = RGBColor(0x00, 0x00, 0x00)         # Black for body text
        self.TABLE_HEADER_BG = RGBColor(0x17, 0x36, 0x5D)    # Blue background for table headers
        self.TABLE_HEADER_TEXT = RGBColor(255, 255, 255)     # White text for headers
        self.TABLE_BORDER_COLOR = RGBColor(0xBF, 0xBF, 0xBF) # Gray borders
        
        # SPACING
        self.SECTION_SPACING_BEFORE = Pt(12)
        self.SECTION_SPACING_AFTER = Pt(6)
        self.PARAGRAPH_SPACING = Pt(3)
        
        # MARGINS
        self.DOC_MARGIN_TOP = Inches(0.8)
        self.DOC_MARGIN_BOTTOM = Inches(0.8)
        self.DOC_MARGIN_LEFT = Inches(0.75)
        self.DOC_MARGIN_RIGHT = Inches(0.75)
    
    def create_complete_template(self, template_path: Path):
        """Create the complete e-Zest CV template with all sections"""
        try:
            doc = Document()
            
            # Set document margins
            self._set_document_margins(doc)
            
            # Add all sections in order
            self._add_header_section(doc)
            self._add_profile_summary(doc)
            self._add_tools_technologies(doc)
            self._add_work_experience(doc)
            self._add_other_projects(doc)
            self._add_education_details(doc)
            self._add_certifications(doc)
            
            # Save template
            doc.save(template_path)
            self.logger.info(f"Complete e-Zest template created at {template_path}")
            
        except Exception as e:
            self.logger.error(f"Error creating template: {str(e)}")
            raise
    
    def _set_document_margins(self, doc):
        """Set document margins"""
        sections = doc.sections
        for section in sections:
            section.top_margin = self.DOC_MARGIN_TOP
            section.bottom_margin = self.DOC_MARGIN_BOTTOM
            section.left_margin = self.DOC_MARGIN_LEFT
            section.right_margin = self.DOC_MARGIN_RIGHT
    
    def _add_header_section(self, doc):
        """Add candidate name and title section"""
        # Candidate Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('{{ contact_info.name }}')
        name_run.font.name = self.HEADING_FONT
        name_run.font.size = self.NAME_FONT_SIZE
        name_run.font.color.rgb = self.HEADING_COLOR
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.space_after = Pt(0)
        
        # Job Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('({{ contact_info.title if contact_info.title else "Professional Title" }})')
        title_run.font.name = self.BODY_FONT
        title_run.font.size = self.TITLE_FONT_SIZE
        title_run.font.color.rgb = self.BODY_COLOR
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.space_after = self.SECTION_SPACING_AFTER
    
    def _add_profile_summary(self, doc):
        """Add Profile Summary section with bullet points"""
        # Section heading
        self._add_section_heading(doc, 'Profile Summary')
        
        # Summary table with bullets
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        cell = table.cell(0, 0)
        
        # Add Jinja2 template for bullet points (using dash instead of bullet character)
        summary_para = cell.paragraphs[0]
        summary_text = """{% if summary_bullets and summary_bullets|length > 0 %}{% for bullet in summary_bullets %}- {{ bullet }}
{% endfor %}{% else %}{{ summary }}{% endif %}"""
        
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.name = self.BODY_FONT
        summary_run.font.size = self.BODY_FONT_SIZE
        summary_run.font.color.rgb = self.BODY_COLOR
        summary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary_para.space_before = self.PARAGRAPH_SPACING
        summary_para.space_after = self.PARAGRAPH_SPACING
        
        self._add_section_spacing(doc)
    
    def _add_tools_technologies(self, doc):
        """Add Tools and Technologies section with 2-column layout"""
        # Section heading
        self._add_section_heading(doc, 'Tools and Technologies')
        
        # Create 2-column table
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Set column widths (approximately 40% and 60%)
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)
        
        # Add Jinja2 template for skills
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        # Left column - Categories
        left_para = left_cell.paragraphs[0]
        left_text = """{% for skill_row in skills_rows %}{{ skill_row.left }}
{% endfor %}"""
        left_run = left_para.add_run(left_text)
        self._format_body_text(left_run)
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Right column - Skills
        right_para = right_cell.paragraphs[0]
        right_text = """{% for skill_row in skills_rows %}{{ skill_row.right }}
{% endfor %}"""
        right_run = right_para.add_run(right_text)
        self._format_body_text(right_run)
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self._add_section_spacing(doc)
    
    def _add_work_experience(self, doc):
        """Add Relevant Work Experience section with detailed project boxes"""
        # Section heading
        self._add_section_heading(doc, 'Relevant Work Experience')
        
        # Experience template with boxes - each project in its own bordered section
        exp_template = """{% for exp in experience %}
Project #{{ loop.index }} Duration: {{ exp.start_date }} - {{ exp.end_date }}

Project Summary: {{ exp.project_description if exp.project_description else exp.description }}

Technologies: {{ exp.technologies|join(", ") if exp.technologies else "N/A" }}

Role & Responsibilities:
{% for resp in exp.responsibilities %}- {{ resp }}
{% endfor %}

{% endfor %}"""
        
        # Create table for structured layout
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        run = para.add_run(exp_template)
        self._format_body_text(run)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self._add_section_spacing(doc)
    
    def _add_other_projects(self, doc):
        """Add Other Notable Projects section with 3-column table"""
        # Section heading
        self._add_section_heading(doc, 'Other Notable Projects')
        
        # Create 3-column table with header
        table = doc.add_table(rows=4, cols=3)  # Header + 3 data rows
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Header row
        headers = ['Project Name', 'Duration', 'Technology']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            self._set_cell_background(cell, self.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = self.HEADING_FONT
            run.font.size = self.TABLE_HEADER_FONT_SIZE
            run.font.color.rgb = self.TABLE_HEADER_TEXT
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template rows for data (rows 1-3)
        for row_idx in range(1, 4):
            # Project name
            para1 = table.cell(row_idx, 0).paragraphs[0]
            run1 = para1.add_run('{% if other_projects and other_projects|length > ' + str(row_idx-1) + ' %}{{ other_projects[' + str(row_idx-1) + '].name }}{% endif %}')
            self._format_body_text(run1)
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Duration
            para2 = table.cell(row_idx, 1).paragraphs[0]
            run2 = para2.add_run('{% if other_projects and other_projects|length > ' + str(row_idx-1) + ' %}{{ other_projects[' + str(row_idx-1) + '].duration }}{% endif %}')
            self._format_body_text(run2)
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Technology
            para3 = table.cell(row_idx, 2).paragraphs[0]
            run3 = para3.add_run('{% if other_projects and other_projects|length > ' + str(row_idx-1) + ' %}{{ other_projects[' + str(row_idx-1) + '].technologies|join(", ") }}{% endif %}')
            self._format_body_text(run3)
            para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_section_spacing(doc)
    
    def _add_education_details(self, doc):
        """Add Education Details section with 3-column table"""
        # Section heading
        self._add_section_heading(doc, 'Education Details')
        
        # Create 3-column table
        table = doc.add_table(rows=2, cols=3)  # Header + 1 data row
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Header row
        headers = ['Course', 'University / Board', 'Year of Passing']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            self._set_cell_background(cell, self.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = self.HEADING_FONT
            run.font.size = self.TABLE_HEADER_FONT_SIZE
            run.font.color.rgb = self.TABLE_HEADER_TEXT
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template row for education data
        # Course
        para1 = table.cell(1, 0).paragraphs[0]
        run1 = para1.add_run('{% for edu in education %}{{ edu.degree }}{% if not loop.last %}, {% endif %}{% endfor %}')
        self._format_body_text(run1)
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # University
        para2 = table.cell(1, 1).paragraphs[0]
        run2 = para2.add_run('{% for edu in education %}{{ edu.institution }}{% if not loop.last %}, {% endif %}{% endfor %}')
        self._format_body_text(run2)
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Year
        para3 = table.cell(1, 2).paragraphs[0]
        run3 = para3.add_run('{% for edu in education %}{{ edu.graduation_date }}{% if not loop.last %}, {% endif %}{% endfor %}')
        self._format_body_text(run3)
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_section_spacing(doc)
    
    def _add_certifications(self, doc):
        """Add Certifications section with 2-column table"""
        # Section heading
        self._add_section_heading(doc, 'Certifications')
        
        # Create 2-column table
        table = doc.add_table(rows=6, cols=2)  # Header + 5 data rows
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Set column widths
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(5.5)
        
        # Header row
        headers = ['Sr.No', 'University/Board']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            self._set_cell_background(cell, self.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = self.HEADING_FONT
            run.font.size = self.TABLE_HEADER_FONT_SIZE
            run.font.color.rgb = self.TABLE_HEADER_TEXT
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template rows for certifications (rows 1-5)
        for i in range(1, 6):
            # Serial number
            para1 = table.cell(i, 0).paragraphs[0]
            run1 = para1.add_run('{% if certifications_rows and certifications_rows|length > ' + str(i-1) + ' %}{{ certifications_rows[' + str(i-1) + '].sno }}{% endif %}')
            self._format_body_text(run1)
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Certification name
            para2 = table.cell(i, 1).paragraphs[0]
            run2 = para2.add_run('{% if certifications_rows and certifications_rows|length > ' + str(i-1) + ' %}{{ certifications_rows[' + str(i-1) + '].authority }}{% endif %}')
            self._format_body_text(run2)
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # HELPER METHODS
    
    def _add_section_heading(self, doc, text):
        """Add a section heading with consistent formatting"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.HEADING_FONT
        run.font.size = self.HEADING_FONT_SIZE
        run.font.color.rgb = self.HEADING_COLOR
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.space_before = self.SECTION_SPACING_BEFORE
        para.space_after = self.SECTION_SPACING_AFTER
    
    def _add_section_spacing(self, doc):
        """Add spacing between sections"""
        para = doc.add_paragraph()
        para.space_after = self.SECTION_SPACING_AFTER
    
    def _format_body_text(self, run):
        """Apply consistent body text formatting"""
        run.font.name = self.BODY_FONT
        run.font.size = self.BODY_FONT_SIZE
        run.font.color.rgb = self.BODY_COLOR
    
    def _set_table_borders(self, table):
        """Set consistent table borders"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:color'), f'{self.TABLE_BORDER_COLOR.rgb:06x}')
                    tcBorders.append(border)
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), f'{color.rgb:06x}')
        tcPr.append(shd)


# USAGE EXAMPLE
def create_ezest_template():
    """Create the complete e-Zest template"""
    generator = EZestCVTemplateGenerator()
    template_path = Path("ezest_complete_template.docx")
    generator.create_complete_template(template_path)
    print(f"Template created: {template_path}")

# Integration with your existing TemplateEngine class
def integrate_with_template_engine():
    """
    Add this method to your TemplateEngine class to use the new generator:
    
    def _create_ezest_complete_template(self):
        ezest_template_path = self.templates_dir / "ezest-complete.docx"
        if not ezest_template_path.exists():
            try:
                generator = EZestCVTemplateGenerator()
                generator.create_complete_template(ezest_template_path)
                logging.info("Complete e-Zest template created successfully")
            except Exception as e:
                logging.error(f"Failed to create complete e-Zest template: {str(e)}")
    """
    pass

# Uncomment to create template
# create_ezest_template()