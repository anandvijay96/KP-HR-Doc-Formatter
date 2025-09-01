from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
import logging


class EZestCVTemplateGenerator:
    """
    Complete e-Zest CV Template Generator with global styling variables
    Based on Rahul Shrivastav CV format with all sections
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # GLOBAL STYLING VARIABLES - Change these to apply across entire template
        self.HEADING_FONT = 'Calibri'          # Main section headings font
        self.BODY_FONT = 'Segoe UI'            # Body text font
        self.NAME_FONT_SIZE = Pt(16)           # Candidate name size
        self.TITLE_FONT_SIZE = Pt(12)          # Job title size
        self.HEADING_FONT_SIZE = Pt(14)        # Section headings size
        self.BODY_FONT_SIZE = Pt(10)           # Regular body text size
        self.TABLE_HEADER_FONT_SIZE = Pt(11)   # Table headers size
        
        # COLORS
        self.HEADING_COLOR = RGBColor(0x17, 0x36, 0x5D)      # Dark blue for headings
        self.BODY_COLOR = RGBColor(0x00, 0x00, 0x00)         # Black for body text
        self.TABLE_HEADER_BG = RGBColor(0x17, 0x36, 0x5D)    # Blue background for table headers
        self.TABLE_HEADER_TEXT = RGBColor(255, 255, 255)     # White text for headers
        self.TABLE_BORDER_COLOR = RGBColor(0xBF, 0xBF, 0xBF) # Gray borders
        
        # SPACING
        self.SECTION_SPACING_BEFORE = Pt(12)
        self.SECTION_SPACING_AFTER = Pt(6)
        self.PARAGRAPH_SPACING = Pt(3)
        
        # MARGINS
        self.DOC_MARGIN_TOP = Inches(0.8)
        self.DOC_MARGIN_BOTTOM = Inches(0.8)
        self.DOC_MARGIN_LEFT = Inches(0.75)
        self.DOC_MARGIN_RIGHT = Inches(0.75)
    
    def create_complete_template(self, template_path: Path):
        """Create the complete e-Zest CV template with all sections"""
        try:
            doc = Document()
            
            # Set document margins
            self._set_document_margins(doc)
            
            # Add all sections in order
            self._add_header_section(doc)
            self._add_profile_summary(doc)
            self._add_tools_technologies(doc)
            self._add_work_experience(doc)
            self._add_other_projects(doc)
            self._add_education_details(doc)
            self._add_certifications(doc)
            
            # Save template
            doc.save(template_path)
            self.logger.info(f"Complete e-Zest template created at {template_path}")
            
        except Exception as e:
            self.logger.error(f"Error creating template: {str(e)}")
            raise
    
    def _set_document_margins(self, doc):
        """Set document margins"""
        sections = doc.sections
        for section in sections:
            section.top_margin = self.DOC_MARGIN_TOP
            section.bottom_margin = self.DOC_MARGIN_BOTTOM
            section.left_margin = self.DOC_MARGIN_LEFT
            section.right_margin = self.DOC_MARGIN_RIGHT
    
    def _add_header_section(self, doc):
    # Candidate Name
        name_para = doc.add_paragraph()
        run = name_para.add_run('{{ contact_info.name }}')
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.color.rgb = self.HEADING_COLOR
        run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.space_after = Pt(0)

        # Job Title
        title_para = doc.add_paragraph()
        run = title_para.add_run('({{ contact_info.title|default("Professional Title") }})')
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.space_after = Pt(6)

        # Horizontal rule (thin gray line)
        p = doc.add_paragraph()._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'BFBFBF')
        pBdr.append(bottom)
        pPr.append(pBdr)

    
    def _add_profile_summary(self, doc):
        """Add Profile Summary section with bullet points"""
        # Section heading
        self._add_section_heading(doc, 'Profile Summary')
        
        # Summary table with bullets
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        cell = table.cell(0, 0)
        
        # Add Jinja2 template for bullet points (using dash instead of bullet character)
        summary_para = cell.paragraphs[0]
        summary_text = """{% if summary_bullets and summary_bullets|length > 0 %}{% for bullet in summary_bullets %}○ {{ bullet }}
{% endfor %}{% else %}{{ summary }}{% endif %}"""
        
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.name = self.BODY_FONT
        summary_run.font.size = self.BODY_FONT_SIZE
        # Avoid explicit RGB assignment for cross-environment compatibility
        summary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary_para.space_before = self.PARAGRAPH_SPACING
        summary_para.space_after = self.PARAGRAPH_SPACING
        
        self._add_section_spacing(doc)
    
    def _add_tools_technologies(self, doc):
        """Add Tools and Technologies section with 2-column layout"""
        # Section heading
        self._add_section_heading(doc, 'Tools and Technologies')

        # Create table with a header row + one template row
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)

        # Set column widths (approximately 40% and 60%)
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)

        # Header row text used as detection marker
        table.cell(0, 0).text = 'Category'
        table.cell(0, 1).text = 'Skills'

        # Template row with simple placeholders (no loops)
        table.cell(1, 0).text = '{{ skill_row.left }}'
        table.cell(1, 1).text = '{{ skill_row.right }}'

        self._add_section_spacing(doc)
    
    def _add_work_experience(self, doc):
        """Add Relevant Work Experience section with detailed project boxes"""
        # Section heading
        self._add_section_heading(doc, 'Relevant Work Experience')
        
        # Experience template with boxes - each project in its own bordered section
        exp_template = """{% for exp in experience %}
Project #{{ loop.index }} Duration: {{ exp.start_date }} - {{ exp.end_date }}

Project Summary: {{ exp.project_description if exp.project_description else exp.description }}

Technologies: {{ exp.technologies|join(", ") if exp.technologies else "N/A" }}

Role & Responsibilities:
{% for resp in exp.responsibilities %}- {{ resp }}
{% endfor %}

{% endfor %}"""
        
        # Create table for structured layout
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        run = para.add_run(exp_template)
        self._format_body_text(run)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self._add_section_spacing(doc)
    
    def _add_other_projects(self, doc):
        """Add Other Notable Projects section with 3-column table"""
        # Section heading
        self._add_section_heading(doc, 'Other Notable Projects')
        
        # Create 3-column table with header
        table = doc.add_table(rows=4, cols=3)  # Header + 3 data rows
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Header row
        headers = ['Project Name', 'Duration', 'Technology']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            self._set_cell_background(cell, self.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = self.HEADING_FONT
            run.font.size = self.TABLE_HEADER_FONT_SIZE
            # Avoid explicit RGB assignment for cross-environment compatibility
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template rows for data (rows 1-3)
        for row_idx in range(1, 4):
            # Project name
            para1 = table.cell(row_idx, 0).paragraphs[0]
            run1 = para1.add_run('{% if other_projects and other_projects|length > ' + str(row_idx-1) + ' %}{{ other_projects[' + str(row_idx-1) + '].name }}{% endif %}')
            self._format_body_text(run1)
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Duration
            para2 = table.cell(row_idx, 1).paragraphs[0]
            run2 = para2.add_run('{% if other_projects and other_projects|length > ' + str(row_idx-1) + ' %}{{ other_projects[' + str(row_idx-1) + '].duration }}{% endif %}')
            self._format_body_text(run2)
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Technology
            para3 = table.cell(row_idx, 2).paragraphs[0]
            run3 = para3.add_run('{% if other_projects and other_projects|length > ' + str(row_idx-1) + ' %}{{ other_projects[' + str(row_idx-1) + '].technologies|join(", ") }}{% endif %}')
            self._format_body_text(run3)
            para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_section_spacing(doc)
    
    def _add_education_details(self, doc):
        """Add Education Details section with 3-column table"""
        # Section heading
        self._add_section_heading(doc, 'Education Details')
        
        # Create 3-column table
        table = doc.add_table(rows=2, cols=3)  # Header + 1 data row
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Header row
        headers = ['Course', 'University / Board', 'Year of Passing']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            self._set_cell_background(cell, self.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = self.HEADING_FONT
            run.font.size = self.TABLE_HEADER_FONT_SIZE
            run.font.color.rgb = self.TABLE_HEADER_TEXT
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template row for education data
        # Course
        para1 = table.cell(1, 0).paragraphs[0]
        run1 = para1.add_run('{% for edu in education %}{{ edu.degree }}{% if not loop.last %}, {% endif %}{% endfor %}')
        self._format_body_text(run1)
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # University
        para2 = table.cell(1, 1).paragraphs[0]
        run2 = para2.add_run('{% for edu in education %}{{ edu.institution }}{% if not loop.last %}, {% endif %}{% endfor %}')
        self._format_body_text(run2)
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Year
        para3 = table.cell(1, 2).paragraphs[0]
        run3 = para3.add_run('{% for edu in education %}{{ edu.graduation_date }}{% if not loop.last %}, {% endif %}{% endfor %}')
        self._format_body_text(run3)
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_section_spacing(doc)
    
    def _add_certifications(self, doc):
        """Add Certifications section with 2-column table"""
        # Section heading
        self._add_section_heading(doc, 'Certifications')
        
        # Create 2-column table
        table = doc.add_table(rows=6, cols=2)  # Header + 5 data rows
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        
        # Set column widths
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(5.5)
        
        # Header row
        headers = ['Sr.No', 'University/Board']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            self._set_cell_background(cell, self.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = self.HEADING_FONT
            run.font.size = self.TABLE_HEADER_FONT_SIZE
            run.font.color.rgb = self.TABLE_HEADER_TEXT
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template rows for certifications (rows 1-5)
        for i in range(1, 6):
            # Serial number
            para1 = table.cell(i, 0).paragraphs[0]
            run1 = para1.add_run('{% if certifications_rows and certifications_rows|length > ' + str(i-1) + ' %}{{ certifications_rows[' + str(i-1) + '].sno }}{% endif %}')
            self._format_body_text(run1)
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Certification name
            para2 = table.cell(i, 1).paragraphs[0]
            run2 = para2.add_run('{% if certifications_rows and certifications_rows|length > ' + str(i-1) + ' %}{{ certifications_rows[' + str(i-1) + '].authority }}{% endif %}')
            self._format_body_text(run2)
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # HELPER METHODS
    
    def _add_section_heading(self, doc, text):
        """Add a section heading with consistent formatting"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.HEADING_FONT
        run.font.size = self.HEADING_FONT_SIZE
        # Avoid explicit RGB assignment for cross-environment compatibility
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.space_before = self.SECTION_SPACING_BEFORE
        para.space_after = self.SECTION_SPACING_AFTER
    
    def _add_section_spacing(self, doc):
        """Add spacing between sections"""
        para = doc.add_paragraph()
        para.space_after = self.SECTION_SPACING_AFTER
    
    def _format_body_text(self, run):
        """Apply consistent body text formatting"""
        run.font.name = self.BODY_FONT
        run.font.size = self.BODY_FONT_SIZE
        # Avoid explicit RGB assignment for cross-environment compatibility
    
    def _set_table_borders(self, table):
        """Set consistent table borders"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:color'), self._rgb_hex(self.TABLE_BORDER_COLOR))
                    tcBorders.append(border)
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), self._rgb_hex(color))
        tcPr.append(shd)

    def _rgb_hex(self, rgb: RGBColor) -> str:
        """Return lowercase hex string RRGGBB for a python-docx RGBColor value."""
        try:
            r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
            return f"{r:02x}{g:02x}{b:02x}"
        except Exception:
            # Fallback: try attribute ._rgb or .rgb if they exist
            try:
                val = getattr(rgb, 'rgb', None) or getattr(rgb, '_rgb', None)
                if val is not None:
                    # val may be bytes or int-like
                    if isinstance(val, (bytes, bytearray)) and len(val) == 3:
                        r, g, b = val[0], val[1], val[2]
                        return f"{r:02x}{g:02x}{b:02x}"
                    try:
                        return f"{int(val):06x}"
                    except Exception:
                        pass
            except Exception:
                pass
            return "000000"

    # ==== Dynamic row helpers and population methods ====
    def _clone_row(self, table, row_idx: int):
        """Clone a row in the table for populating data while preserving formatting."""
        from copy import deepcopy
        tr = table.rows[row_idx]._tr
        new_tr = deepcopy(tr)
        table._tbl.append(new_tr)
        return table.rows[-1]

    def _clear_data_rows(self, table, header_rows: int = 1):
        """Remove all rows after the given number of header rows."""
        while len(table.rows) > header_rows:
            table._tbl.remove(table.rows[-1]._tr)

    def _find_table_by_headers(self, doc, headers: list[str]):
        """Find a table whose first row's cell texts match the provided headers (startswith match)."""
        hdrs_lower = [h.strip().lower() for h in headers]
        for t in doc.tables:
            if len(t.rows) == 0:
                continue
            first = t.rows[0]
            if len(first.cells) < len(headers):
                continue
            row_texts = [first.cells[i].text.strip().lower() for i in range(len(headers))]
            # startswith allows minor variations
            if all(row_texts[i].startswith(hdrs_lower[i]) for i in range(len(headers))):
                return t
        return None

    def _remove_row(self, table, row_idx: int):
        """Remove a row from a python-docx table by index."""
        try:
            row = table.rows[row_idx]
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)
        except Exception:
            # Fallback: clear cell texts to avoid rendering placeholders
            if 0 <= row_idx < len(table.rows):
                for cell in table.rows[row_idx].cells:
                    cell.text = ""

    def populate_skills(self, doc, skills_rows: list[dict]):
        """Populate the Tools and Technologies table using 'Category'/'Skills' headers."""
        table = self._find_table_by_headers(doc, ['Category', 'Skills'])
        if table is None:
            raise ValueError("Skills table not found (headers 'Category'|'Skills')")
        # Keep header row (0) and template row (1)
        self._clear_data_rows(table, header_rows=2)
        template_row_idx = 1
        for row in skills_rows or []:
            new_row = self._clone_row(table, template_row_idx)
            new_row.cells[0].text = str(row.get('left', '')).strip()
            new_row.cells[1].text = str(row.get('right', '')).strip()
        # Remove the template row
        if len(table.rows) > 1:
            self._remove_row(table, template_row_idx)

    def populate_other_projects(self, doc, other_projects: list[dict]):
        """Populate the Other Notable Projects 3-col table by headers."""
        table = self._find_table_by_headers(doc, ['Project Name', 'Duration', 'Technology'])
        if table is None:
            raise ValueError("Other Notable Projects table not found")
        # Ensure there is a template row at index 1
        if len(table.rows) > 1:
            header_rows = 2
            template_row_idx = 1
        else:
            # Add a blank template row
            table.add_row()
            header_rows = 2
            template_row_idx = 1
        # Clear data rows but keep header + template
        self._clear_data_rows(table, header_rows=header_rows)
        for p in other_projects or []:
            new_row = self._clone_row(table, template_row_idx)
            new_row.cells[0].text = str(p.get('name', '')).strip()
            new_row.cells[1].text = str(p.get('duration', '')).strip()
            techs = p.get('technologies') or []
            if isinstance(techs, list):
                techs_text = ', '.join([str(x) for x in techs if str(x).strip()])
            else:
                techs_text = str(techs)
            new_row.cells[2].text = techs_text
        # Remove the template row
        if len(table.rows) > 1:
            self._remove_row(table, template_row_idx)

    def populate_education(self, doc, education: list[dict]):
        """Populate the Education Details 3-col table by headers."""
        table = self._find_table_by_headers(doc, ['Course', 'University / Board', 'Year of Passing'])
        if table is None:
            raise ValueError("Education table not found")
        if len(table.rows) > 1:
            header_rows = 2
            template_row_idx = 1
        else:
            table.add_row()
            header_rows = 2
            template_row_idx = 1
        self._clear_data_rows(table, header_rows=header_rows)
        for e in education or []:
            new_row = self._clone_row(table, template_row_idx)
            new_row.cells[0].text = str(e.get('degree', '')).strip()
            new_row.cells[1].text = str(e.get('institution', '')).strip()
            new_row.cells[2].text = str(e.get('graduation_date', '')).strip()
        # Remove the template row
        if len(table.rows) > 1:
            self._remove_row(table, template_row_idx)

    def populate_certifications(self, doc, cert_rows: list[dict]):
        """Populate the Certifications 2-col table by headers."""
        table = self._find_table_by_headers(doc, ['Sr.No', 'University/Board'])
        if table is None:
            raise ValueError("Certifications table not found")
        if len(table.rows) > 1:
            header_rows = 2
            template_row_idx = 1
        else:
            table.add_row()
            header_rows = 2
            template_row_idx = 1
        self._clear_data_rows(table, header_rows=header_rows)
        for r in cert_rows or []:
            new_row = self._clone_row(table, template_row_idx)
            new_row.cells[0].text = str(r.get('sno', '')).strip()
            new_row.cells[1].text = str(r.get('authority', '')).strip()
        # Remove the template row
        if len(table.rows) > 1:
            self._remove_row(table, template_row_idx)


# USAGE EXAMPLE
def create_ezest_template():
    """Create the complete e-Zest template"""
    generator = EZestCVTemplateGenerator()
    template_path = Path("ezest_complete_template.docx")
    generator.create_complete_template(template_path)
    print(f"Template created: {template_path}")

# Integration with your existing TemplateEngine class
def integrate_with_template_engine():
    """
    Add this method to your TemplateEngine class to use the new generator:
    
    def _create_ezest_complete_template(self):
        ezest_template_path = self.templates_dir / "ezest-complete.docx"
        if not ezest_template_path.exists():
            try:
                generator = EZestCVTemplateGenerator()
                generator.create_complete_template(ezest_template_path)
                logging.info("Complete e-Zest template created successfully")
            except Exception as e:
                logging.error(f"Failed to create complete e-Zest template: {str(e)}")
    """
    pass

# Uncomment to create template
# create_ezest_template()